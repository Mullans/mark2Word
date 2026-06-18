"""Docx emission tests for spacing, dual-align, images, and hyperlinks."""

import base64
import re
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from mark2word import ImageError, build

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"

MINIMAL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


class SpacingTests(unittest.TestCase):
    def test_body_space_after_applies_before_next_paragraph(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "spacing.docx"
            build(
                glob={"body": {"space_after": 8}},
                external={},
                md_body="First\nSecond\n",
                out_path=out,
            )
            doc = Document(out)
            self.assertEqual(doc.paragraphs[0].paragraph_format.space_after.pt, 0)
            self.assertEqual(doc.paragraphs[1].paragraph_format.space_before.pt, 8)

    def test_list_space_between_applies_between_items(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "list-spacing.docx"
            build(
                glob={"list": {"space_between": 5}},
                external={},
                md_body="- one\n- two\n",
                out_path=out,
            )
            doc = Document(out)
            bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
            self.assertEqual(bullets[1].paragraph_format.space_before.pt, 5)

    def test_last_paragraph_keeps_space_after(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "tail-spacing.docx"
            build(
                glob={"body": {"space_after": 6}},
                external={},
                md_body="Only\n",
                out_path=out,
            )
            doc = Document(out)
            self.assertEqual(doc.paragraphs[0].paragraph_format.space_after.pt, 6)


class DualAlignTests(unittest.TestCase):
    def test_dual_align_renders_tab_between_sides(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "dual.docx"
            build(
                glob={},
                external={},
                md_body="Left side || Right side\n",
                out_path=out,
            )
            doc = Document(out)
            para = doc.paragraphs[0]
            self.assertIn("Left side", para.text)
            self.assertIn("Right side", para.text)
            tabs = para._p.findall(f".//{W}tab")
            self.assertTrue(tabs, "expected a tab element between dual-aligned sides")
            p_pr = para._p.find(f"{W}pPr")
            self.assertIsNotNone(p_pr)
            tab_stops = p_pr.findall(f"{W}tabs/{W}tab")
            self.assertEqual(len(tab_stops), 1)
            self.assertEqual(tab_stops[0].get(f"{W}val"), "right")


class HyperlinkTests(unittest.TestCase):
    def test_hyperlink_emits_relationship_and_element(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "link.docx"
            build(
                glob={},
                external={},
                md_body="See [Example](https://example.com) here\n",
                out_path=out,
            )
            with zipfile.ZipFile(out) as z:
                doc_xml = z.read("word/document.xml")
                rels_xml = z.read("word/_rels/document.xml.rels")
            self.assertIn(b"w:hyperlink", doc_xml)
            self.assertIn(b"https://example.com", rels_xml)


class ImageTests(unittest.TestCase):
    def test_embedded_image_renders_in_docx(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            img = root / "pixel.png"
            img.write_bytes(MINIMAL_PNG)
            out = root / "img.docx"
            build(
                glob={},
                external={},
                md_body=f"![swatch]({img.name})\n",
                out_path=out,
                md_path=root / "doc.md",
            )
            doc = Document(out)
            self.assertEqual(len(doc.inline_shapes), 1)

    def test_image_alt_text_in_docx_metadata(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            img = root / "pixel.png"
            img.write_bytes(MINIMAL_PNG)
            out = root / "img-alt.docx"
            build(
                glob={"image": {"alt_mode": "doc"}},
                external={},
                md_body=f"![Swatch description]({img.name})\n",
                out_path=out,
                md_path=root / "doc.md",
            )
            with zipfile.ZipFile(out) as z:
                doc_xml = z.read("word/document.xml").decode("utf-8")
            self.assertIn('descr="Swatch description"', doc_xml)
            body_text = re.sub(r"<w:drawing>.*?</w:drawing>", "", doc_xml, flags=re.DOTALL)
            self.assertNotIn("Swatch description", body_text)

    def test_missing_image_raises_image_error(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            md = root / "doc.md"
            md.write_text("![missing](nope.png)\n", encoding="utf-8")
            out = root / "out.docx"
            with self.assertRaises(ImageError) as ctx:
                build(
                    glob={},
                    external={},
                    md_body="![missing](nope.png)\n",
                    out_path=out,
                    md_path=md,
                )
            self.assertIn("image not found", str(ctx.exception))
            self.assertEqual(ctx.exception.line_no, 1)


if __name__ == "__main__":
    unittest.main()
