"""Tests for blockquote, hr, pagebreak, internal links, theme fidelity, CLI."""

import base64
import io
import unittest
import zipfile
from contextlib import redirect_stderr
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from mark2word import build, parse_blocks
from mark2word.slug import slugify
from mark2word.cli import EXIT_FAILURE, EXIT_SUCCESS, main
from mark2word.theme import split_frontmatter

MINIMAL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


class MarkdownFeatureTests(unittest.TestCase):
    def test_blockquote_parses(self):
        blocks = parse_blocks("> quoted text\n> second line\n")
        self.assertEqual(blocks[0]["type"], "blockquote")
        self.assertIn("quoted text", blocks[0]["text"])

    def test_horizontal_rule_in_body(self):
        front, body = split_frontmatter("---\ntitle: Doc\n---\n---\n")
        self.assertEqual(front["title"], "Doc")
        blocks = parse_blocks(body)
        self.assertEqual(blocks[0]["type"], "hr")

    def test_pagebreak_comment(self):
        blocks = parse_blocks("Before\n<!-- pagebreak -->\nAfter\n")
        self.assertEqual(blocks[1]["type"], "pagebreak")

    def test_blockquote_table_layout(self):
        from mark2word.emit import _blockquote_table_layout
        from mark2word.theme import border_width_twips, length_to_twips

        content = 10_368
        style = {
            "padding": {"left": "0.1in", "right": "0.1in"},
            "border_left": {"size": "3pt", "color": "D6D6D6"},
        }
        indent, width = _blockquote_table_layout(style, content)
        cell_left = length_to_twips("0.1in")
        cell_right = length_to_twips("0.1in")
        border = border_width_twips(style["border_left"])
        self.assertEqual(indent, cell_left + border)
        self.assertEqual(width, content - indent + cell_right)

    def test_blockquote_indent_left_reduces_width(self):
        from mark2word.emit import _blockquote_table_layout

        base_style = {"padding": {"left": "0.1in", "right": "0.1in"}}
        indent_a, width_a = _blockquote_table_layout(base_style, 9000)
        indent_b, width_b = _blockquote_table_layout(
            {**base_style, "indent_left": "12pt"}, 9000,
        )
        self.assertGreater(indent_b, indent_a)
        self.assertLess(width_b, width_a)
        self.assertEqual(width_a - width_b, indent_b - indent_a)

    def test_fill_none_skips_shading(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "no-fill.docx"
            build(
                glob={"code_block": {"fill": "none"}, "code_inline": {"fill": "none"}},
                external={},
                md_body="`x`\n```\ny\n```\n",
                out_path=out,
            )
            with zipfile.ZipFile(out) as z:
                xml = z.read("word/document.xml").decode("utf-8")
            self.assertNotIn('w:fill="F5F5F5"', xml)

    def test_region_top_level_props_apply_to_body(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "region-body.docx"
            build(
                glob={
                    "$pullquote": {
                        "font": "Georgia",
                        "size": 13,
                        "color": "C00000",
                        "italic": True,
                        "body": {"align": "center"},
                    },
                },
                external={},
                md_body="<!-- region: pullquote -->\nRegion line\n<!-- /region -->\n",
                out_path=out,
            )
            doc = Document(out)
            para = doc.paragraphs[0]
            self.assertEqual(para.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            run = para.runs[0]
            self.assertEqual(run.font.name, "Georgia")
            self.assertEqual(run.font.size.pt, 13)
            self.assertTrue(run.italic)
            self.assertEqual(str(run.font.color.rgb), "C00000")

    def test_blockquote_space_before_on_preceding_paragraph(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "bq-space.docx"
            build(
                glob={"blockquote": {"space_before": 12, "space_after": 0}},
                external={},
                md_body="Body line\n\n> quoted\n",
                out_path=out,
            )
            doc = Document(out)
            body = doc.paragraphs[0]
            self.assertEqual(body.paragraph_format.space_after.pt, 12)
            quote_cell = doc.tables[0].rows[0].cells[0].paragraphs[0]
            self.assertEqual(quote_cell.paragraph_format.space_before.pt, 0)

    def test_blockquote_renders_in_docx(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "bq.docx"
            build(
                glob={"blockquote": {"italic": True}},
                external={},
                md_body="> A quoted line\n",
                out_path=out,
            )
            with zipfile.ZipFile(out) as z:
                xml = z.read("word/document.xml").decode("utf-8")
            self.assertIn("quoted", xml.lower())
            self.assertIn("<w:tbl>", xml)
            self.assertIn('w:fill="F0F0F0"', xml)
            self.assertIn("<w:tblInd", xml)
            self.assertIn("<w:tblW", xml)

    def test_inline_code_run_shading(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "inline-code.docx"
            build(
                glob={},
                external={},
                md_body="Use `inline` here\n",
                out_path=out,
            )
            with zipfile.ZipFile(out) as z:
                xml = z.read("word/document.xml").decode("utf-8")
            self.assertIn('w:fill="F5F5F5"', xml)
            self.assertIn("inline", xml)

    def test_code_block_defaults_to_monospace(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "mono.docx"
            build(glob={}, external={}, md_body="```\nx\n```\n", out_path=out)
            with zipfile.ZipFile(out) as z:
                styles = z.read("word/styles.xml").decode("utf-8")
            self.assertIn("Consolas", styles)

    def test_hr_renders_border(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "hr.docx"
            build(glob={}, external={}, md_body="---\n", out_path=out)
            with zipfile.ZipFile(out) as z:
                xml = z.read("word/document.xml")
            self.assertIn(b"w:pBdr", xml)

    def test_internal_link_to_heading(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "links.docx"
            build(
                glob={},
                external={},
                md_body="# My Section\n\nSee [link](#my-section).\n",
                out_path=out,
            )
            with zipfile.ZipFile(out) as z:
                doc_xml = z.read("word/document.xml").decode("utf-8")
            self.assertIn('w:anchor="my-section"', doc_xml)

    def test_unknown_internal_link_raises(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "bad.docx"
            with self.assertRaises(Exception):
                build(
                    glob={},
                    external={},
                    md_body="[nope](#missing)\n",
                    out_path=out,
                )

    def test_bold_hyperlink(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "bold-link.docx"
            build(
                glob={},
                external={},
                md_body="[**bold link**](https://example.com)\n",
                out_path=out,
            )
            with zipfile.ZipFile(out) as z:
                xml = z.read("word/document.xml")
            self.assertIn(b"w:hyperlink", xml)
            self.assertIn(b"w:b", xml)

    def test_indent_right(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "ind.docx"
            build(
                glob={"body": {"indent_right": "24pt"}},
                external={},
                md_body="Indented\n",
                out_path=out,
            )
            doc = Document(out)
            self.assertEqual(doc.paragraphs[0].paragraph_format.right_indent.pt, 24)

    def test_table_spacing(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "tspace.docx"
            build(
                glob={"table": {"space_before": 8, "space_after": 6}},
                external={},
                md_body="Above\n| A | B |\n| - | - |\n| 1 | 2 |\nBelow\n",
                out_path=out,
            )
            doc = Document(out)
            table = doc.tables[0]
            self.assertEqual(
                table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before.pt, 8
            )

    def test_image_width_and_caption(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            img = root / "pic.png"
            img.write_bytes(MINIMAL_PNG)
            out = root / "img.docx"
            build(
                glob={"image": {"width": "50pt", "alt_mode": "caption", "align": "center"}},
                external={},
                md_body=f"![Swatch label]({img.name})\n",
                out_path=out,
                md_path=root / "doc.md",
            )
            doc = Document(out)
            self.assertEqual(len(doc.inline_shapes), 1)
            self.assertTrue(any("Swatch" in p.text for p in doc.paragraphs))

    def test_code_lang_theme_override(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "code.docx"
            build(
                glob={"code": {"font": "Consolas", "langs": {"python": {"color": "0000FF"}}}},
                external={},
                md_body="```python\nx = 1\n```\n",
                out_path=out,
            )
            doc = Document(out)
            run = doc.paragraphs[0].runs[0]
            self.assertEqual(str(run.font.color.rgb), "0000FF")

    def test_code_fill_applies_paragraph_shading(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "code-fill.docx"
            build(
                glob={},
                external={},
                md_body="```\nline one\nline two\n```\n",
                out_path=out,
            )
            with zipfile.ZipFile(out) as z:
                doc_xml = z.read("word/document.xml").decode("utf-8")
            self.assertIn('w:fill="F5F5F5"', doc_xml)
            self.assertIn("Mark2wordCodeBlock", doc_xml)
            self.assertNotRegex(doc_xml, r"<w:tbl>.*?F5F5F5")

    def test_page_header_footer_from_theme(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "chrome.docx"
            build(
                glob={
                    "title": "My Doc",
                    "page": {"footer": "My Doc || Page {page}"},
                },
                external={},
                md_body="# Title\n",
                out_path=out,
            )
            with zipfile.ZipFile(out) as z:
                footer = z.read("word/footer1.xml").decode("utf-8")
                styles = z.read("word/styles.xml").decode("utf-8")
            self.assertIn("My Doc", footer)
            self.assertIn("fldChar", footer)
            self.assertIn('w:val="clear"', footer)
            self.assertIn('w:val="right"', footer)
            footer_style = styles[styles.find('w:styleId="Footer"'):]
            footer_style = footer_style[: footer_style.find("</w:style>")]
            self.assertNotIn('w:val="center"', footer_style)

    def test_slugify(self):
        self.assertEqual(slugify("Hello World!"), "hello-world")


class CliFeatureTests(unittest.TestCase):
    def test_version_flag(self):
        with self.assertRaises(SystemExit) as ctx:
            main(["--version"])
        self.assertEqual(ctx.exception.code, 0)

    def test_check_theme_ok(self):
        theme = Path("skills/mark2word/assets/base-theme.yaml")
        if not theme.exists():
            self.skipTest("skill assets not present")
        code = main(["--check-theme", str(theme), "--theme-dir", "skills/mark2word/assets"])
        self.assertEqual(code, EXIT_SUCCESS)

    def test_check_theme_bad_page_size(self):
        with TemporaryDirectory() as tmp:
            bad = Path(tmp) / "bad.yaml"
            bad.write_text("page:\n  size: legal\n", encoding="utf-8")
            stderr = io.StringIO()
            with redirect_stderr(stderr):
                code = main(["--check-theme", str(bad)])
            self.assertEqual(code, EXIT_FAILURE)
            self.assertIn("unknown page size", stderr.getvalue())


if __name__ == "__main__":
    unittest.main()
