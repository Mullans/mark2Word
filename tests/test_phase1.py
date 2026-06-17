import unittest
import warnings
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from mark2word import (
    Block,
    FrontmatterError,
    RegionError,
    Resolver,
    ThemeError,
    build,
    deep_merge,
    load_theme,
    parse_blocks,
    parse_inline,
    parse_to_ast,
    register_block_parser,
    split_frontmatter,
)
from mark2word.cli import main, resolve_output_path, validate_document
from mark2word.parser import split_dual_align
from mark2word.theme import discover_theme_dirs


class FrontmatterTests(unittest.TestCase):
    def test_lf_frontmatter(self):
        front, body = split_frontmatter("---\ntitle: One\n---\n# Body\n")
        self.assertEqual(front, {"title": "One"})
        self.assertEqual(body, "# Body\n")

    def test_crlf_frontmatter(self):
        text = "---\r\ntitle: Two\r\n---\r\n# Body\r\n"
        front, body = split_frontmatter(text)
        self.assertEqual(front, {"title": "Two"})
        self.assertEqual(body, "# Body\r\n")

    def test_bom_frontmatter(self):
        text = "\ufeff---\ntitle: Three\n---\nBody\n"
        front, body = split_frontmatter(text)
        self.assertEqual(front, {"title": "Three"})
        self.assertEqual(body, "Body\n")

    def test_no_frontmatter(self):
        front, body = split_frontmatter("# Just markdown\n")
        self.assertEqual(front, {})
        self.assertEqual(body, "# Just markdown\n")

    def test_body_may_contain_triple_dashes(self):
        text = "---\ntitle: Doc\n---\n# Title\n\n---\nnot frontmatter\n"
        front, body = split_frontmatter(text)
        self.assertEqual(front, {"title": "Doc"})
        self.assertIn("---\nnot frontmatter", body)

    def test_rejects_non_mapping_frontmatter(self):
        with self.assertRaises(FrontmatterError):
            split_frontmatter("---\n- one\n- two\n---\n")

    def test_empty_frontmatter_block(self):
        front, body = split_frontmatter("---\n---\nBody\n")
        self.assertEqual(front, {})
        self.assertEqual(body, "Body\n")


class DeepMergeTests(unittest.TestCase):
    def test_page_partial_override(self):
        merged = deep_merge(
            {"size": "letter", "margin": {"top": "0.5in", "left": "0.7in"}},
            {"margin": {"top": "1in"}},
        )
        self.assertEqual(merged["size"], "letter")
        self.assertEqual(merged["margin"]["top"], "1in")
        self.assertEqual(merged["margin"]["left"], "0.7in")


class RegionTests(unittest.TestCase):
    def test_full_line_region_open(self):
        blocks = parse_blocks("<!-- region: main -->\n# Title\n<!-- /region -->")
        self.assertEqual(blocks[0]["region"], ["main"])
        self.assertEqual(blocks[0]["type"], "h1")

    def test_inline_region_comment_is_body_text(self):
        blocks = parse_blocks("Note <!-- region: main -->\n")
        self.assertEqual(len(blocks), 1)
        self.assertEqual(blocks[0]["type"], "body")

    def test_unclosed_region_fails(self):
        with self.assertRaises(RegionError) as ctx:
            parse_blocks("<!-- region: main -->\n# Title\n")
        self.assertIn("unclosed region: main", str(ctx.exception))

    def test_extra_close_fails_with_line_number(self):
        with self.assertRaises(RegionError) as ctx:
            parse_blocks("<!-- /region -->\n")
        self.assertEqual(ctx.exception.line_no, 1)


class BuildPageMergeTests(unittest.TestCase):
    def test_frontmatter_margin_overrides_theme_only_for_set_keys(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            out = root / "out.docx"
            build(
                glob={"page": {"margin": {"top": "1in"}}},
                external={"page": {"size": "letter", "margin": {"top": "0.5in", "left": "0.8in"}}},
                md_body="# Hello\n",
                out_path=out,
            )
            doc = Document(out)
            sec = doc.sections[0]
            self.assertEqual(sec.top_margin.inches, 1.0)
            self.assertEqual(sec.left_margin.inches, 0.8)


class Tier2Tests(unittest.TestCase):
    def test_parse_inline_link_and_underscore(self):
        runs = parse_inline("_hi_ and [site](https://example.com)")
        self.assertTrue(any(r.italic and r.text == "hi" for r in runs))
        self.assertTrue(any(r.url == "https://example.com" for r in runs))

    def test_parse_inline_code(self):
        runs = parse_inline("use `print()` here")
        self.assertTrue(any(r.code and r.text == "print()" for r in runs))

    def test_nested_list_levels(self):
        blocks = parse_blocks("- top\n  - nested\n  - nested2\n- top2\n")
        self.assertEqual([b["list_level"] for b in blocks], [0, 1, 1, 0])

    def test_nested_ordered_list_levels(self):
        blocks = parse_blocks("1. one\n   1. nested\n2. two\n")
        self.assertEqual([b["list_level"] for b in blocks], [0, 1, 0])

    def test_nested_bullets_use_ilvl_in_docx(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "nested.docx"
            build(
                glob={},
                external={},
                md_body="- top\n  - nested\n  - nested2\n- top2\n",
                out_path=out,
            )
            doc = Document(out)
            bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]

            def ilvl(paragraph):
                num_pr = paragraph._p.pPr.numPr
                return num_pr.ilvl.val if num_pr is not None and num_pr.ilvl is not None else None

            self.assertEqual([ilvl(p) for p in bullets], [0, 1, 1, 0])

    def test_nested_list_indent_step(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "nested-indent.docx"
            build(
                glob={},
                external={"list": {"indent_left": "12pt", "indent_hanging": "9pt", "indent_step": "18pt"}},
                md_body="- top\n  - nested\n    - deep\n",
                out_path=out,
            )
            doc = Document(out)
            bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
            indents = [p.paragraph_format.left_indent.pt for p in bullets]
            self.assertAlmostEqual(indents[0], 12.0, places=1)
            self.assertAlmostEqual(indents[1], 30.0, places=1)
            self.assertAlmostEqual(indents[2], 48.0, places=1)

    def test_list_level_style_override(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "level-color.docx"
            build(
                glob={},
                external={
                    "list": {
                        "indent_left": "10pt",
                        "indent_step": "10pt",
                        "levels": {2: {"color": "FF0000"}},
                    }
                },
                md_body="- a\n  - b\n    - c\n",
                out_path=out,
            )
            doc = Document(out)
            bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
            deep_run = bullets[2].runs[0]
            self.assertEqual(str(deep_run.font.color.rgb), "FF0000")

    def test_nested_ordered_uses_ilvl_in_docx(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "nested-ol.docx"
            build(
                glob={},
                external={},
                md_body="1. one\n   1. nested\n2. two\n",
                out_path=out,
            )
            doc = Document(out)
            numbers = [p for p in doc.paragraphs if p.style.name == "List Number"]

            def ilvl(paragraph):
                num_pr = paragraph._p.pPr.numPr
                return num_pr.ilvl.val if num_pr is not None and num_pr.ilvl is not None else None

            self.assertEqual([ilvl(p) for p in numbers], [0, 1, 0])

    def test_nested_ordered_restarts_numbering_per_level(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "nested-ol-numbers.docx"
            build(
                glob={},
                external={},
                md_body="1. item one\n2. item two\n  1. nested\n3. item three\n",
                out_path=out,
            )
            doc = Document(out)
            numbers = [p for p in doc.paragraphs if p.style.name == "List Number"]

            def num_info(paragraph):
                num_pr = paragraph._p.pPr.numPr
                num_id = num_pr.numId.val if num_pr is not None and num_pr.numId is not None else None
                ilvl = num_pr.ilvl.val if num_pr is not None and num_pr.ilvl is not None else None
                return num_id, ilvl

            info = [num_info(p) for p in numbers]
            self.assertEqual([i[1] for i in info], [0, 0, 1, 0])
            top_num_id = info[0][0]
            self.assertEqual(top_num_id, info[1][0])
            self.assertEqual(top_num_id, info[2][0])
            self.assertEqual(top_num_id, info[3][0])

    def test_ordered_list_uses_decimal_multilevel_abstract(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "ol-abstract.docx"
            build(
                glob={},
                external={},
                md_body="1. one\n  1. nested\n",
                out_path=out,
            )
            import zipfile
            import xml.etree.ElementTree as ET

            W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            with zipfile.ZipFile(out) as z:
                num_xml = z.read("word/numbering.xml").decode("utf-8")
            nroot = ET.fromstring(num_xml)
            decimal_abstracts = []
            for ab in nroot.findall("w:abstractNum", ns):
                lvls = ab.findall("w:lvl", ns)
                if not lvls:
                    continue
                fmts = [
                    lvl.find("w:numFmt", ns).get(f"{W}val")
                    for lvl in lvls
                    if lvl.find("w:numFmt", ns) is not None
                ]
                if fmts and all(fmt == "decimal" for fmt in fmts):
                    aid = ab.get(f"{W}abstractNumId")
                    decimal_abstracts.append(aid)
            self.assertTrue(decimal_abstracts)

    def test_fenced_code_block(self):
        blocks = parse_blocks("```py\nprint(1)\n```\n")
        self.assertEqual(blocks[0]["type"], "code")
        self.assertEqual(blocks[0]["code_lang"], "py")

    def test_a4_page_size(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "a4.docx"
            build(
                glob={"page": {"size": "a4"}},
                external={},
                md_body="Body\n",
                out_path=out,
            )
            doc = Document(out)
            self.assertAlmostEqual(doc.sections[0].page_width.inches, 8.27, places=2)

    def test_heading_uses_word_heading_style(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "h.docx"
            build(
                glob={"h1": {"size": 16}},
                external={},
                md_body="# Title\n",
                out_path=out,
            )
            doc = Document(out)
            self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")
            self.assertEqual(doc.paragraphs[0].runs[0].font.size.pt, 16)

    def test_discover_theme_dirs(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            themes = root / ".mark2word" / "themes"
            themes.mkdir(parents=True)
            md = root / "doc.md"
            md.write_text("# x\n", encoding="utf-8")
            self.assertIn(themes.resolve(), [p.resolve() for p in discover_theme_dirs(md)])

    def test_check_validation(self):
        with TemporaryDirectory() as tmp:
            md = Path(tmp) / "ok.md"
            md.write_text("---\nfont: Calibri\n---\n# Hi\n", encoding="utf-8")
            validate_document(md, [])

    def test_ordered_list_warning_on_skipped_marker(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "list.docx"
            with warnings.catch_warnings(record=True) as caught:
                warnings.simplefilter("always")
                build(
                    glob={},
                    external={},
                    md_body="1. one\n3. three\n",
                    out_path=out,
                    verbose=True,
                )
            self.assertTrue(any("ignored" in str(w.message) for w in caught))

    def test_split_dual_align_skips_backticks(self):
        self.assertIsNone(split_dual_align("use `left || right` inline"))
        self.assertEqual(
            split_dual_align("Alpha || Omega"),
            ("Alpha", "Omega"),
        )
        self.assertIsNone(split_dual_align(r"literal \|\| bars"))

    def test_ordered_list_after_bullets_uses_numbers(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "lists.docx"
            build(
                glob={},
                external={},
                md_body="- bullet\n1. first\n2. second\n",
                out_path=out,
            )
            doc = Document(out)
            styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
            self.assertIn("List Bullet", styles)
            self.assertIn("List Number", styles)

    def test_table_header_text_renders(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "table.docx"
            build(
                glob={},
                external={"th": {"bold": True, "color": "FFFFFF", "fill": "2B579A"}},
                md_body="| Col A | Col B |\n| - | - |\n| one | two |\n",
                out_path=out,
            )
            doc = Document(out)
            self.assertEqual(doc.tables[0].cell(0, 0).text, "Col A")
            self.assertEqual(doc.tables[0].cell(0, 1).text, "Col B")


class Tier3Tests(unittest.TestCase):
    def test_theme_extends_chain(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            base = root / "base.yaml"
            child = root / "child.yaml"
            md = root / "doc.md"
            base.write_text("font: Base\nbody: { space_after: 2pt }\n", encoding="utf-8")
            child.write_text("extends: base.yaml\nsize: 12\n", encoding="utf-8")
            md.write_text("---\nextends: child.yaml\n---\n", encoding="utf-8")
            external, glob = load_theme({"extends": "child.yaml"}, md)
            self.assertEqual(external["font"], "Base")
            self.assertEqual(external["size"], 12)

    def test_theme_extends_cycle(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            a = root / "a.yaml"
            b = root / "b.yaml"
            a.write_text("extends: b.yaml\n", encoding="utf-8")
            b.write_text("extends: a.yaml\n", encoding="utf-8")
            with self.assertRaises(ThemeError):
                load_theme({"extends": "a.yaml"}, root / "doc.md")

    def test_table_block(self):
        blocks = parse_blocks("| a | b |\n| - | - |\n| 1 | 2 |\n")
        self.assertEqual(blocks[0]["type"], "table")
        self.assertEqual(blocks[0]["rows"][1], ["1", "2"])

    def test_table_emits_in_docx(self):
        with TemporaryDirectory() as tmp:
            out = Path(tmp) / "table.docx"
            build(
                glob={},
                external={},
                md_body="| Name | Val |\n| - | - |\n| x | 1 |\n",
                out_path=out,
            )
            doc = Document(out)
            self.assertEqual(len(doc.tables), 1)
            self.assertEqual(doc.tables[0].cell(1, 0).text, "x")

    def test_batch_converts_multiple_files(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            one = root / "one.md"
            two = root / "two.md"
            out_dir = root / "docx"
            one.write_text("# One\n", encoding="utf-8")
            two.write_text("# Two\n", encoding="utf-8")
            from mark2word.cli import main

            main(["-i", str(one), "-i", str(two), "-o", str(out_dir)])
            self.assertTrue((out_dir / "one.docx").exists())
            self.assertTrue((out_dir / "two.docx").exists())


class Tier4Tests(unittest.TestCase):
    def test_ast_roundtrip(self):
        block = Block(type="body", text="hello", region=["r"])
        self.assertEqual(Block.from_dict(block.to_dict()).text, "hello")

    def test_resolver_is_memoized(self):
        resolver = Resolver({}, {})
        resolver.resolve("body")
        resolver.resolve("body")
        self.assertEqual(len(resolver._cache), 1)
        self.assertEqual(resolver.resolve("body")["font"], "Calibri")

    def test_plugin_block_parser(self):
        def parse_banner(stripped, line_no, ctx):
            if stripped.startswith("!!!"):
                return {"type": "body", "text": stripped[3:].strip(), "region": list(ctx["region_stack"])}
            return None

        register_block_parser(parse_banner, prepend=True)
        blocks = parse_blocks("!!! Banner text\n")
        self.assertEqual(blocks[0]["text"], "Banner text")


if __name__ == "__main__":
    unittest.main()
