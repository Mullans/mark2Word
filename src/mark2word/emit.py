"""Word document emission."""

from __future__ import annotations

import warnings
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt

from mark2word.ast import Block
from mark2word.errors import ImageError, ParseError, ThemeError
from mark2word.lists import apply_list_numbering
from mark2word.oxml_helpers import (
    add_bookmark,
    add_dual_align_tab,
    add_hyperlink as oxml_add_hyperlink,
    add_page_break,
    clear_table_borders,
    prepare_dual_align_paragraph,
    set_cell_borders,
    set_cell_margins,
    set_cell_shading,
    set_paragraph_border_bottom,
    set_paragraph_shading,
    set_run_shading,
    set_style_paragraph_shading,
    set_image_alt,
    set_table_borders,
    set_table_layout,
)
from mark2word.parser import parse_inline, parse_to_ast, split_dual_align
from mark2word.plugins import block_emitter
from mark2word.sections import configure_page_chrome
from mark2word.slug import unique_slug
from mark2word.theme import (
    HEADINGS,
    PAGE_SIZES,
    NumberingConfig,
    Resolver,
    border_enabled,
    border_width_twips,
    deep_merge,
    fill_enabled,
    hex_color,
    length_to_twips,
    pt_value,
    to_length,
)

MARK2WORD_BODY = "Mark2word Body"
MARK2WORD_CODE_BLOCK = "Mark2word Code Block"

ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


@dataclass
class FlowItem:
    kind: str
    element: Any
    block: dict[str, Any]
    style: dict[str, Any]


def _resolve_image_path(block: dict[str, Any], md_path: Path | None) -> Path:
    image_path = Path(block["image_path"])
    if not image_path.is_absolute() and md_path is not None:
        image_path = md_path.parent / image_path
    return image_path


def _picture_width(inline, style: dict[str, Any]) -> None:
    if style.get("width") is not None:
        inline.width = int(to_length(style["width"]))
    if style.get("max_width") is not None:
        max_emu = int(to_length(style["max_width"]))
        if inline.width > max_emu:
            ratio = max_emu / inline.width
            inline.width = max_emu
            inline.height = int(inline.height * ratio)


def _insert_image(paragraph, image_path: Path, style: dict[str, Any], alt: str, *, line_no: int | None) -> None:
    if not image_path.is_file():
        raise ImageError(f"image not found: {image_path}", line_no)
    try:
        run = paragraph.add_run()
        inline = run.add_picture(str(image_path))
        _picture_width(inline, style)
        alt_mode = str(style.get("alt_mode", "doc")).lower()
        if alt and alt_mode in {"doc", "both"}:
            set_image_alt(run, alt)
    except OSError as exc:
        raise ImageError(f"cannot read image: {image_path}", line_no) from exc


def apply_paragraph_style(paragraph, style: dict[str, Any]) -> None:
    pf = paragraph.paragraph_format
    if style.get("align") in ALIGN:
        paragraph.alignment = ALIGN[style["align"]]
    line = style.get("line")
    if line is not None:
        if isinstance(line, (int, float)):
            pf.line_spacing = float(line)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        else:
            pf.line_spacing = to_length(line)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if style.get("indent_left") is not None:
        pf.left_indent = to_length(style["indent_left"])
    if style.get("indent_right") is not None:
        pf.right_indent = to_length(style["indent_right"])
    if style.get("indent_first_line") is not None:
        pf.first_line_indent = to_length(style["indent_first_line"])
    if style.get("indent_hanging") is not None:
        pf.first_line_indent = -to_length(style["indent_hanging"])
    if style.get("border_bottom"):
        set_paragraph_border_bottom(paragraph, style["border_bottom"])
    if fill_enabled(style.get("fill")):
        set_paragraph_shading(paragraph, style["fill"])


def apply_run_style(run, style: dict[str, Any], *, inline_code: bool = False) -> None:
    if style.get("font"):
        run.font.name = style["font"]
    if style.get("size") is not None:
        run.font.size = Pt(style["size"])
    if style.get("color"):
        run.font.color.rgb = hex_color(style["color"])
    if style.get("bold") is not None:
        run.bold = style["bold"]
    if style.get("italic") is not None:
        run.italic = style["italic"]
    if inline_code and fill_enabled(style.get("fill")):
        set_run_shading(run, style["fill"])


def add_runs(
    paragraph,
    segment: str,
    style: dict[str, Any],
    *,
    inline_code_style: dict[str, Any] | None = None,
    code_style: dict[str, Any] | None = None,
    anchors: set[str] | None = None,
    inherit_paragraph_style: bool = False,
) -> None:
    for run in parse_inline(segment):
        if run.code and inline_code_style is not None:
            active = inline_code_style
            use_inline = True
        elif run.code and code_style is not None:
            active = code_style
            use_inline = False
        else:
            active = style
            use_inline = False
        if run.url:
            if run.url.startswith("#"):
                anchor = run.url[1:]
                if anchors is not None and anchor not in anchors:
                    raise ParseError(f"unknown internal link anchor: #{anchor}")
                oxml_add_hyperlink(
                    paragraph, run.text, run.url, active,
                    bold=run.bold, italic=run.italic,
                )
            else:
                oxml_add_hyperlink(
                    paragraph, run.text, run.url, active,
                    bold=run.bold, italic=run.italic,
                )
            continue
        if inherit_paragraph_style and not run.code and not run.bold and not run.italic:
            paragraph.add_run(run.text)
            continue
        r = paragraph.add_run(run.text)
        apply_run_style(r, active, inline_code=use_inline)
        if run.bold:
            r.bold = True
        if run.italic:
            r.italic = True


def _ensure_paragraph_style(doc: Document, name: str, base_name: str, resolved: dict[str, Any]):
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH, True)
        style.base_style = doc.styles[base_name]
    style.font.name = resolved.get("font", style.font.name)
    if resolved.get("size") is not None:
        style.font.size = Pt(resolved["size"])
    if resolved.get("color"):
        style.font.color.rgb = hex_color(resolved["color"])
    if resolved.get("bold") is not None:
        style.font.bold = resolved["bold"]
    if resolved.get("italic") is not None:
        style.font.italic = resolved["italic"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    if fill_enabled(resolved.get("fill")):
        set_style_paragraph_shading(style, resolved["fill"])


def configure_document_styles(doc: Document, resolver: Resolver) -> None:
    base = resolver.resolve("body")
    normal = doc.styles["Normal"]
    normal.font.name = base.get("font", "Calibri")
    normal.font.size = Pt(base.get("size", 11))
    if base.get("color"):
        normal.font.color.rgb = hex_color(base["color"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    _ensure_paragraph_style(doc, MARK2WORD_BODY, "Normal", base)
    _ensure_paragraph_style(
        doc, MARK2WORD_CODE_BLOCK, "Normal", resolver.resolve("code_block"),
    )

    for level in range(1, 7):
        element = f"h{level}"
        style = resolver.resolve(element)
        heading = doc.styles[f"Heading {level}"]
        if style.get("font"):
            heading.font.name = style["font"]
        if style.get("size") is not None:
            heading.font.size = Pt(style["size"])
        if style.get("color"):
            heading.font.color.rgb = hex_color(style["color"])
        if style.get("bold") is not None:
            heading.font.bold = style["bold"]
        if style.get("italic") is not None:
            heading.font.italic = style["italic"]


def _heading_style_name(block_type: str) -> str | None:
    if block_type in HEADINGS:
        return f"Heading {block_type[1]}"
    return None


def _warn_list_number_mismatch(expected: int, actual: int, line_no: int | None) -> None:
    where = f"line {line_no}" if line_no else "document"
    warnings.warn(
        f"ordered list number {actual} at {where} ignored; continuing from {expected}",
        stacklevel=2,
    )


def _assign_heading_anchors(blocks: list[Block]) -> set[str]:
    used: set[str] = set()
    for block in blocks:
        if block.type.startswith("h") and block.text:
            block.anchor = unique_slug(block.text, used)
    return used


def emit_paragraph(
    doc,
    block: dict[str, Any],
    style: dict[str, Any],
    content_width_twips: int,
    *,
    prev_list_paragraph=None,
    list_run_num_id: int | None = None,
    expected_ol_number: int | None = None,
    numbering_config: NumberingConfig | None = None,
    verbose: bool = False,
    md_path: Path | None = None,
    code_style: dict[str, Any] | None = None,
    inline_code_style: dict[str, Any] | None = None,
    anchors: set[str] | None = None,
    bookmark_id: int | None = None,
):
    block_type = block["type"]
    custom = block_emitter(block_type)
    if custom is not None:
        p = custom(doc, block, style, {
            "content_width_twips": content_width_twips,
            "md_path": md_path,
            "code_style": code_style,
        })
        return p, None

    if block_type == "pagebreak":
        p = doc.add_paragraph()
        add_page_break(p)
        return p, None

    list_num_id = None
    if block_type == "list":
        ordered = block.get("ordered", False)
        p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        level = int(block.get("list_level", 0))
        start = block.get("list_number") if ordered else None
        if ordered and expected_ol_number is not None and start is not None and start != expected_ol_number:
            if verbose:
                _warn_list_number_mismatch(expected_ol_number, start, block.get("line_no"))
        list_num_id = apply_list_numbering(
            p, doc, ordered=ordered, start=start, level=level,
            prev=prev_list_paragraph, run_num_id=list_run_num_id,
            numbering_config=numbering_config,
        )
    elif block_type == "code":
        p = doc.add_paragraph(style=MARK2WORD_CODE_BLOCK)
    elif block_type == "hr":
        p = doc.add_paragraph(style=MARK2WORD_BODY)
        if "border_bottom" not in style:
            style = {**style, "border_bottom": {"size": "0.5pt", "color": "999999"}}
    elif heading := _heading_style_name(block_type):
        p = doc.add_paragraph(style=heading)
    else:
        p = doc.add_paragraph(style=MARK2WORD_BODY)

    apply_paragraph_style(p, style)
    pf = p.paragraph_format

    if block.get("anchor") and bookmark_id is not None:
        add_bookmark(p, bookmark_id, block["anchor"])

    if block_type == "image":
        image_path = _resolve_image_path(block, md_path)
        image_style = {k: v for k, v in style.items() if k in style}
        alt = block.get("image_alt", "")
        _insert_image(p, image_path, image_style, alt, line_no=block.get("line_no"))
        alt_mode = str(style.get("alt_mode", "doc")).lower()
        if alt and alt_mode in {"caption", "both"}:
            cap = doc.add_paragraph(style=MARK2WORD_BODY)
            cap.alignment = ALIGN.get(style.get("align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
            cap.add_run(alt)
            return p, None
        return p, None

    if block_type == "code":
        lines = block.get("text", "").splitlines() or [""]
        if len(lines) == 1:
            apply_run_style(p.add_run(lines[0]), code_style or style)
        else:
            for line in lines:
                apply_run_style(p.add_run(line + "\n"), code_style or style)
        return p, None

    # Outside regions, body runs inherit Mark2word Body; inside a region, apply
    # resolved run props (region top-level font/size/color + nested body keys).
    inherit = block_type == "body" and not block.get("region")
    parts = split_dual_align(block.get("text", ""))
    if parts is not None:
        left, right = parts
        prepare_dual_align_paragraph(p, content_width_twips)
        add_runs(p, left, style, code_style=code_style, inline_code_style=inline_code_style, anchors=anchors, inherit_paragraph_style=inherit)
        add_dual_align_tab(p)
        add_runs(p, right, style, code_style=code_style, inline_code_style=inline_code_style, anchors=anchors, inherit_paragraph_style=inherit)
    else:
        add_runs(
            p, block.get("text", ""), style, code_style=code_style, inline_code_style=inline_code_style, anchors=anchors,
            inherit_paragraph_style=inherit,
        )
    return p, list_num_id


def _blockquote_cell_padding(style: dict[str, Any]) -> dict[str, Any]:
    padding = style.get("padding")
    return dict(padding) if isinstance(padding, dict) else {}


def _blockquote_table_layout(
    style: dict[str, Any],
    content_width_twips: int,
) -> tuple[int, int]:
    """
    Table indent and width so the blockquote spans the text column.

    indent = theme indent_left + cell padding-left + left border width
    width  = content_width - indent + cell padding-right
    """
    padding = _blockquote_cell_padding(style)
    cell_left = length_to_twips(padding.get("left", "0.1in"))
    cell_right = length_to_twips(padding.get("right", "0.1in"))
    user_indent = length_to_twips(style.get("indent_left"))
    border_left = style.get("border_left")
    border_twips = border_width_twips(border_left) if border_enabled(border_left) else 0
    table_indent = user_indent + cell_left + border_twips
    table_width = content_width_twips - table_indent + cell_right
    return table_indent, max(0, table_width)


def _blockquote_border_spec(style: dict[str, Any]) -> dict[str, Any]:
    spec: dict[str, Any] = {}
    for edge in ("left", "right"):
        cfg = style.get(f"border_{edge}")
        if border_enabled(cfg):
            spec[edge] = cfg
    return spec


def emit_blockquote(
    doc: Document,
    block: dict[str, Any],
    style: dict[str, Any],
    content_width_twips: int,
    *,
    inline_code_style: dict[str, Any] | None = None,
    anchors: set[str] | None = None,
    bookmark_id: int | None = None,
) -> Any:
    """Blockquote as a single-cell table aligned to the document text column."""
    table = doc.add_table(rows=1, cols=1)
    clear_table_borders(table)
    table_indent, table_width = _blockquote_table_layout(style, content_width_twips)
    set_table_layout(table, indent_twips=table_indent, width_twips=table_width)
    cell = table.rows[0].cells[0]
    if fill_enabled(style.get("fill")):
        set_cell_shading(cell, style["fill"])
    set_cell_borders(cell, _blockquote_border_spec(style))
    padding = _blockquote_cell_padding(style)
    if padding:
        set_cell_margins(cell, padding)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if block.get("anchor") and bookmark_id is not None:
        add_bookmark(p, bookmark_id, block["anchor"])
    lines = block.get("text", "").splitlines() or [""]
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run("\n")
        add_runs(
            p, line, style,
            inline_code_style=inline_code_style,
            anchors=anchors,
            inherit_paragraph_style=True,
        )
    return table


def emit_table(doc, block: dict[str, Any], resolver: Resolver) -> Any:
    rows = block.get("rows", [])
    if not rows:
        return None
    region = tuple(block.get("region", []))
    table_style = resolver.resolve("table", region)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    border = table_style.get("border")
    if isinstance(border, dict):
        set_table_borders(table, border)
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            element = "th" if r_idx == 0 else "td"
            cell_style = resolver.resolve(element, region)
            merged = {**table_style, **cell_style}
            if fill_enabled(merged.get("fill")):
                set_cell_shading(cell, merged["fill"])
            padding = merged.get("padding")
            if isinstance(padding, dict):
                set_cell_margins(cell, padding)
            p = cell.paragraphs[0]
            add_runs(p, cell_text, merged)
    return table


def _is_blockquote_table(item: FlowItem) -> bool:
    return item.kind == "table" and item.block.get("type") == "blockquote"


def realize_flow_spacing(items: list[FlowItem]) -> None:
    for idx, item in enumerate(items):
        before = pt_value(item.style.get("space_before", 0))
        if idx == 0:
            # No preceding block: must place leading space on the item itself.
            # Blockquote tables keep space out of the shaded cell when possible;
            # at document start there is no predecessor, so this one case stays
            # on the first in-cell paragraph.
            if before > 0:
                if _is_blockquote_table(item):
                    _apply_blockquote_cell_space_before(item, before)
                else:
                    _apply_space_before(item, before)
            continue
        prev = items[idx - 1]
        same_list = (
            item.kind == "paragraph"
            and prev.kind == "paragraph"
            and item.block.get("type") == prev.block.get("type") == "list"
        )
        if same_list and ("space_between" in prev.style or "space_between" in item.style):
            gap = pt_value(item.style.get("space_between", prev.style.get("space_between", 0)))
        else:
            gap = max(pt_value(prev.style.get("space_after", 0)), before)
        if _is_blockquote_table(item):
            # Move blockquote space_before to the preceding block's space_after so
            # the gap sits above the table, not inside the shaded cell.
            _apply_space_after(prev, gap)
        else:
            _apply_space_after(prev, 0)
            _apply_space_before(item, gap)
    if items:
        _apply_space_after(items[-1], pt_value(items[-1].style.get("space_after", 0)))


def _apply_space_before(item: FlowItem, pt: float) -> None:
    if _is_blockquote_table(item):
        return
    if item.kind == "paragraph":
        item.element.paragraph_format.space_before = Pt(pt)
    elif item.kind == "table" and item.element.rows:
        item.element.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(pt)


def _apply_blockquote_cell_space_before(item: FlowItem, pt: float) -> None:
    if item.kind == "table" and item.element.rows:
        item.element.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(pt)


def _apply_space_after(item: FlowItem, pt: float) -> None:
    if item.kind == "paragraph":
        item.element.paragraph_format.space_after = Pt(pt)
    elif item.kind == "table" and item.element.rows:
        last_row = item.element.rows[-1]
        last_row.cells[-1].paragraphs[0].paragraph_format.space_after = Pt(pt)


def emit_from_ast(
    doc: Document,
    blocks: list[Block],
    resolver: Resolver,
    *,
    content_width_twips: int,
    verbose: bool = False,
    md_path: Path | None = None,
) -> list[FlowItem]:
    flow: list[FlowItem] = []
    prev_list_p = None
    prev_list_ordered: bool | None = None
    list_run_num_id: int | None = None
    ol_expected: dict[int, int] = {}
    numbering_configs: dict[tuple[tuple[str, ...], bool], NumberingConfig] = {}
    anchors = _assign_heading_anchors(blocks)
    bookmark_seq = 0

    for block in blocks:
        data = block.to_dict()
        block_type = data["type"]
        region = tuple(data.get("region", []))

        if block_type == "table":
            table_style = resolver.resolve("table", region)
            table = emit_table(doc, data, resolver)
            if table is not None:
                flow.append(FlowItem("table", table, data, table_style))
            prev_list_p = None
            prev_list_ordered = None
            list_run_num_id = None
            ol_expected.clear()
            continue

        is_ordered = bool(data.get("ordered"))
        numbering_config = None
        if block_type == "list":
            cache_key = (region, is_ordered)
            if cache_key not in numbering_configs:
                numbering_configs[cache_key] = resolver.resolve_numbering_config(
                    region, ordered=is_ordered,
                )
            numbering_config = numbering_configs[cache_key]
            style = resolver.resolve_list_style(
                int(data.get("list_level", 0)), region, ordered=is_ordered,
            )
        elif block_type == "code":
            style = resolver.resolve("code_block", region)
        elif block_type == "image":
            style = resolver.resolve("image", region)
        else:
            style = resolver.resolve(block_type, region)

        inline_code_style = resolver.resolve_code_inline_style(region)
        code_style = None
        if block_type == "code":
            code_style = resolver.resolve_code_block_style(data.get("code_lang", ""), region)

        if block_type == "blockquote":
            bookmark_id = None
            if data.get("anchor"):
                bookmark_id = bookmark_seq
                bookmark_seq += 1
            table = emit_blockquote(
                doc, data, style, content_width_twips,
                inline_code_style=inline_code_style,
                anchors=anchors,
                bookmark_id=bookmark_id,
            )
            flow.append(FlowItem("table", table, data, style))
            prev_list_p = None
            prev_list_ordered = None
            list_run_num_id = None
            ol_expected.clear()
            continue

        expected = None
        same_list_run = (
            block_type == "list"
            and prev_list_p is not None
            and prev_list_ordered == is_ordered
        )
        run_num_id = list_run_num_id if same_list_run else None

        if block_type == "list" and is_ordered:
            level = int(data.get("list_level", 0))
            if not same_list_run:
                ol_expected.clear()
            if level not in ol_expected:
                ol_expected[level] = data.get("list_number") or 1
            expected = ol_expected[level]
            ol_expected[level] = expected + 1

        if block_type == "code" and fill_enabled(style.get("fill")):
            lines = data.get("text", "").splitlines() or [""]
            if len(lines) > 1:
                bookmark_id = None
                if data.get("anchor"):
                    bookmark_id = bookmark_seq
                    bookmark_seq += 1
                for i, line in enumerate(lines):
                    line_style = {**style}
                    if i > 0:
                        line_style["space_before"] = 0
                    if i < len(lines) - 1:
                        line_style["space_after"] = 0
                    line_block = {**data, "text": line}
                    p, _ = emit_paragraph(
                        doc, line_block, line_style, content_width_twips,
                        verbose=verbose,
                        md_path=md_path,
                        code_style=code_style,
                        inline_code_style=inline_code_style,
                        anchors=anchors,
                        bookmark_id=bookmark_id if i == 0 else None,
                    )
                    flow.append(FlowItem("paragraph", p, data, line_style))
                prev_list_p = None
                prev_list_ordered = None
                list_run_num_id = None
                ol_expected.clear()
                continue

        bookmark_id = None
        if data.get("anchor"):
            bookmark_id = bookmark_seq
            bookmark_seq += 1

        p, new_num_id = emit_paragraph(
            doc, data, style, content_width_twips,
            prev_list_paragraph=prev_list_p if same_list_run else None,
            list_run_num_id=run_num_id,
            expected_ol_number=expected,
            numbering_config=numbering_config,
            verbose=verbose,
            md_path=md_path,
            code_style=code_style,
            inline_code_style=inline_code_style,
            anchors=anchors,
            bookmark_id=bookmark_id,
        )
        if block_type == "list":
            prev_list_p = p
            prev_list_ordered = is_ordered
            list_run_num_id = new_num_id
        else:
            prev_list_p = None
            prev_list_ordered = None
            list_run_num_id = None
            ol_expected.clear()
        flow.append(FlowItem("paragraph", p, data, style))
    return flow


def build(
    glob: dict[str, Any],
    external: dict[str, Any],
    md_body: str,
    out_path: Path,
    *,
    verbose: bool = False,
    md_path: Path | None = None,
) -> None:
    resolver = Resolver(external, glob)
    doc = Document()

    page = deep_merge(external.get("page", {}), glob.get("page", {}))
    sec = doc.sections[0]
    size_key = str(page.get("size", "letter")).lower()
    if size_key not in PAGE_SIZES:
        raise ThemeError(f"unknown page size: {page.get('size')!r}")
    sec.page_width, sec.page_height = PAGE_SIZES[size_key]
    margin = page.get("margin", {})
    sec.top_margin = to_length(margin.get("top", "0.5in"))
    sec.bottom_margin = to_length(margin.get("bottom", "0.5in"))
    sec.left_margin = to_length(margin.get("left", "0.7in"))
    sec.right_margin = to_length(margin.get("right", "0.7in"))
    content_emu = int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)
    content_width_twips = int(content_emu / 635)

    doc_title = str(glob.get("title") or external.get("title") or "")

    configure_document_styles(doc, resolver)
    blocks = parse_to_ast(md_body)
    if not doc_title:
        for block in blocks:
            if block.type == "h1" and block.text:
                doc_title = block.text
                break

    configure_page_chrome(
        doc, page, doc_title=doc_title, content_width_twips=content_width_twips,
    )

    flow = emit_from_ast(
        doc, blocks, resolver,
        content_width_twips=content_width_twips,
        verbose=verbose, md_path=md_path,
    )
    realize_flow_spacing(flow)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")
