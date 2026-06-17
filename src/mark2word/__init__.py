import argparse
import re
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


DEFAULTS = {
    "font": "Calibri",
    "size": 11, "color": "000000",
    "body": {},
    "list": {"indent_left": "18pt", "indent_hanging": "18pt"},
    "heading": {"bold": True},
}

PROP_KEYS = {
    "font", "size", "color", "bold", "italic", "align", "line",
    "space_before", "space_between", "space_after",
    "indent_left", "indent_hanging", "indent_first_line", "border_bottom",
}
HEADINGS = {f"h{i}" for i in range(1, 7)}
TEXT_ELEMENTS = {"body", "list"}
ELEMENT_KEYS = HEADINGS | TEXT_ELEMENTS | {"text", "heading"}

ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def to_length(v):
    """Convert pt/in/bare to docx lengths (pt/bare -> Pt, in -> Inches)"""
    if isinstance(v, (int, float)):
        return Pt(v)
    s = str(v).strip().lower()
    if s.endswith("pt"):
        return Pt(float(s[:-2]))
    if s.endswith("in"):
        return Inches(float(s[:-2]))
    return Pt(float(s))


def hex_color(v):
    """Get color from color hex string"""
    return RGBColor.from_string(str(v).lstrip("#"))


def pt_value(v):
    """Convert pt/in to float."""
    if v is None:
        return 0.0
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().lower()
    if s.endswith("pt"):
        return float(s[:-2])
    if s.endswith("in"):
        return float(s[:-2]) * 72.0
    return float(s)


def split_frontmatter(text):
    """Extract and load the YAML frontmatter"""
    m = re.match(r"^---\n(.*?)\n---\n?(.*)$", text, re.S)
    if not m:
        return {}, text
    front = yaml.safe_load(m.group(1)) or {}
    return front, m.group(2)


def _theme_candidates(ext: Path, md_path: Path, theme_dirs=None):
    if ext.is_absolute():
        yield ext
        return
    yield md_path.parent / ext
    for theme_dir in theme_dirs or []:
        yield Path(theme_dir).expanduser() / ext


def load_theme(
    front: dict[str, Any],
    md_path: Path,
    theme_dirs=None,
) -> tuple[dict[str, Any], dict[str, Any]]:
    """Load the `extends` path if present.

    Relative theme paths are searched in the markdown file's folder first, then
    in any explicit theme directories supplied by the caller.
    """
    external = {}
    front = dict(front)
    ext = front.pop("extends", None)
    if ext is not None:
        checked = [path.resolve() for path in _theme_candidates(Path(ext), md_path, theme_dirs)]
        ext_path = next((path for path in checked if path.exists()), None)
        if ext_path is None:
            searched = ", ".join(str(path) for path in checked)
            raise FileNotFoundError(f"extends references missing theme: {ext}; searched: {searched}")
        external = yaml.safe_load(ext_path.read_text(encoding="utf-8")) or {}
    return external, front


# --------------------------------------------------------------------------- #
# Cascade resolution
# --------------------------------------------------------------------------- #
def _bare(layer):
    return {k: v for k, v in layer.items() if k in PROP_KEYS}


def _key_order(element):
    if element in HEADINGS:
        return ["heading", element]
    if element in TEXT_ELEMENTS:
        return ["text", element]
    return [element]


def _layer_contribution(layer, element):
    """Merge a single layer's relevant keys for `element`: bare -> umbrella -> specific."""
    out = dict(_bare(layer))
    for key in _key_order(element):
        sub = layer.get(key)
        if isinstance(sub, dict):
            out.update({k: v for k, v in sub.items() if k in PROP_KEYS})
    return out


class Resolver:
    def __init__(self, external, glob):
        # source layers, low -> high, that carry a *global* scope
        self.global_layers = [DEFAULTS, external, glob]
        # the same external/global layers also hold (nested) $region scopes
        self.region_sources = [external, glob]

    @staticmethod
    def _descend(layer, path):
        """Follow $-prefixed keys down a region path; None if any level is absent."""
        node = layer
        for name in path:
            if not isinstance(node, dict):
                return None
            node = node.get("$" + name)
        return node if isinstance(node, dict) else None

    def resolve(self, element, region_path=None):
        style = {}
        # scope: global  (defaults < external < frontmatter)
        for layer in self.global_layers:
            style.update(_layer_contribution(layer, element))
        # scope: regions, outermost -> innermost. A deeper region always beats a
        # shallower one; within one depth, frontmatter beats external.
        for depth in range(1, len(region_path or []) + 1):
            prefix = region_path[:depth]
            for src in self.region_sources:
                block = self._descend(src, prefix)
                if block is not None:
                    style.update(_layer_contribution(block, element))
        return style


RE_REGION_OPEN = re.compile(r"<!--\s*region:\s*([A-Za-z0-9_-]+)\s*-->")
RE_REGION_CLOSE = re.compile(r"<!--\s*/region\s*-->")
RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
RE_UL = re.compile(r"^[-*]\s+(.*)$")
RE_OL = re.compile(r"^\d+\.\s+(.*)$")


def parse_blocks(md):
    blocks = []
    region_stack = []
    for raw in md.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        mo = RE_REGION_OPEN.search(stripped)
        if mo:
            region_stack.append(mo.group(1))
            continue
        if RE_REGION_CLOSE.search(stripped):
            if region_stack:
                region_stack.pop()
            continue
        region = list(region_stack)   # full path, outermost -> innermost

        mh = RE_HEADING.match(stripped)
        if mh:
            blocks.append({"type": f"h{len(mh.group(1))}", "text": mh.group(2).strip(),
                           "region": region})
            continue
        mu = RE_UL.match(stripped)
        if mu:
            blocks.append({"type": "list", "ordered": False, "text": mu.group(1).strip(),
                           "region": region})
            continue
        mol = RE_OL.match(stripped)
        if mol:
            blocks.append({"type": "list", "ordered": True, "text": mol.group(1).strip(),
                           "region": region})
            continue
        blocks.append({"type": "body", "text": stripped, "region": region})
    return blocks


def parse_inline(text):
    """-> list of (text, bold, italic) runs."""
    runs = []
    buf = []
    bold = italic = False
    i, n = 0, len(text)

    def flush():
        if buf:
            runs.append(("".join(buf), bold, italic))
            buf.clear()

    while i < n:
        c = text[i]
        if c == "\\" and i + 1 < n:
            buf.append(text[i + 1])
            i += 2
            continue
        if c == "*":
            run = len(text[i:]) - len(text[i:].lstrip("*"))
            stars = min(run, 3)
            flush()
            if stars == 3:
                bold = not bold
                italic = not italic
            elif stars == 2:
                bold = not bold
            else:
                italic = not italic
            i += stars
            continue
        buf.append(c); i += 1
    flush()
    return runs


def apply_run_style(run, style):
    if style.get("font"):
        run.font.name = style["font"]
    if style.get("size") is not None:
        run.font.size = Pt(style["size"])
    if style.get("color"):
        run.font.color.rgb = hex_color(style["color"])
    if style.get("bold") is not None:
        run.bold = style["bold"]
    if style.get("italic") is not None:
        run.italic = style["italic"]


def add_runs(paragraph, segment, style):
    for txt, b, it in parse_inline(segment):
        r = paragraph.add_run(txt)
        apply_run_style(r, style)
        # inline markup overrides the element default for that property
        if b:
            r.bold = True
        if it:
            r.italic = True


def set_border_bottom(paragraph, spec):
    sz = max(1, round(float(str(spec.get("size", "0.5pt")).rstrip("pt")) * 8))  # pt -> eighths
    color = str(spec.get("color", "000000")).lstrip("#")
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="2" w:color="{color}"/>'
        f'</w:pBdr>'))


def emit_paragraph(doc, block, style, content_width_twips):
    ordered = block.get("ordered", False)
    if block["type"] == "list":
        p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
    else:
        p = doc.add_paragraph()

    pf = p.paragraph_format

    if style.get("align") in ALIGN:
        p.alignment = ALIGN[style["align"]]

    # dual alignment:  left || right
    parts = re.split(r"\s*\|\|\s*", block["text"], maxsplit=1)
    if len(parts) == 2:
        pf.tab_stops.add_tab_stop(Twips(content_width_twips), WD_TAB_ALIGNMENT.RIGHT)
        add_runs(p, parts[0], style)
        p.add_run().add_tab()
        add_runs(p, parts[1], style)
    else:
        add_runs(p, block["text"], style)

    # line spacing: number => multiple, "Npt" => exact
    line = style.get("line")
    if line is not None:
        if isinstance(line, (int, float)):
            pf.line_spacing = float(line)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        else:
            pf.line_spacing = to_length(line)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    # indents
    if style.get("indent_left") is not None:
        pf.left_indent = to_length(style["indent_left"])
    if style.get("indent_first_line") is not None:
        pf.first_line_indent = to_length(style["indent_first_line"])
    if style.get("indent_hanging") is not None:        # hanging = negative first-line
        pf.first_line_indent = -to_length(style["indent_hanging"])

    if style.get("border_bottom"):
        set_border_bottom(p, style["border_bottom"])

    return p


def realize_spacing(doc_paragraphs):
    """Collapse every inter-paragraph boundary so exactly one side carries the gap.

    Same-type adjacent items use that type's `space_between`; a type change uses
    max(upper.space_after, lower.space_before). One side is always zeroed, so the
    rendered gap is deterministic regardless of Word's sum/max compatibility flag.
    """
    for idx, (p, blk, st) in enumerate(doc_paragraphs):
        before = pt_value(st.get("space_before", 0))
        if idx == 0:
            p.paragraph_format.space_before = Pt(before)
            continue
        up_p, up_blk, up_st = doc_paragraphs[idx - 1]
        same = (blk["type"] == up_blk["type"] == "list")
        if same and ("space_between" in up_st or "space_between" in st):
            gap = pt_value(st.get("space_between", up_st.get("space_between", 0)))
        else:
            gap = max(pt_value(up_st.get("space_after", 0)), before)
        up_p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(gap)
    if doc_paragraphs:
        last_p, _, last_st = doc_paragraphs[-1]
        last_p.paragraph_format.space_after = Pt(pt_value(last_st.get("space_after", 0)))


def build(front, glob, external, md_body, md_path, out_path):
    resolver = Resolver(external, glob)
    doc = Document()

    page = glob.get("page", external.get("page", {}))
    sec = doc.sections[0]
    if page.get("size", "letter") == "letter":
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    else:
        raise ValueError(f"Unknown page size: `{page.get("size")}`")
    margin = page.get("margin", {})
    sec.top_margin = to_length(margin.get("top", "0.5in"))
    sec.bottom_margin = to_length(margin.get("bottom", "0.5in"))
    sec.left_margin = to_length(margin.get("left", "0.7in"))
    sec.right_margin = to_length(margin.get("right", "0.7in"))
    content_emu = int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)
    content_width_twips = int(content_emu / 635)   # 1 twip = 635 EMU

    # Normal baseline so anything we don't touch still matches defaults
    base = resolver.resolve("body")
    normal = doc.styles["Normal"]
    normal.font.name = base.get("font", "Calibri")
    normal.font.size = Pt(base.get("size", 11))
    if base.get("color"):
        normal.font.color.rgb = hex_color(base["color"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    emitted = []
    for blk in parse_blocks(md_body):
        style = resolver.resolve(blk["type"], blk.get("region"))
        p = emit_paragraph(doc, blk, style, content_width_twips)
        emitted.append((p, blk, style))

    realize_spacing(emitted)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Convert a styled-Markdown dialect into a Word document."
    )
    parser.add_argument(
        "--theme-dir",
        action="append",
        default=[],
        type=Path,
        help="Additional folder to search for relative frontmatter extends paths.",
    )
    parser.add_argument("input_file", type=Path)
    parser.add_argument("output_file", nargs="?", type=Path)
    args = parser.parse_args()

    md_path = args.input_file
    out_path = args.output_file or md_path.with_suffix(".docx")
    text = md_path.read_text(encoding="utf-8")
    front, body = split_frontmatter(text)
    external, glob = load_theme(front, md_path, theme_dirs=args.theme_dir)
    build(front, glob, external, body, md_path, out_path)


if __name__ == "__main__":
    main()
