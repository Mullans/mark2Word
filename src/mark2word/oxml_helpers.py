"""Centralized Word OXML helpers (python-docx public API + minimal private access)."""

from __future__ import annotations

from typing import Any

from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"


def paragraph_element(paragraph) -> Any:
    return paragraph._p


def cell_element(cell) -> Any:
    return cell._tc


def add_run_text(parent, text: str, *, bold: bool = False, italic: bool = False) -> None:
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    if len(r_pr):
        run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.set(qn("xml:space"), "preserve")
    text_elem.text = text
    run.append(text_elem)
    parent.append(run)


def add_hyperlink(
    paragraph,
    text: str,
    url: str,
    style: dict[str, Any],
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """External (http) or internal (#anchor) hyperlink."""
    hyperlink = OxmlElement("w:hyperlink")
    if url.startswith("#"):
        hyperlink.set(qn("w:anchor"), url[1:])
    else:
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if style.get("color"):
        color = OxmlElement("w:color")
        color.set(qn("w:val"), str(style["color"]).lstrip("#"))
        r_pr.append(color)
    if style.get("underline") is not False:
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    if len(r_pr):
        run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.set(qn("xml:space"), "preserve")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph_element(paragraph).append(hyperlink)


def _make_shading_element(fill: str) -> OxmlElement:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), str(fill).lstrip("#"))
    return shading


def set_paragraph_shading(paragraph, fill: str) -> None:
    """Paragraph background (Word Home → Shading): w:pPr/w:shd."""
    p_pr = paragraph_element(paragraph).get_or_add_pPr()
    existing = p_pr.find(qn("w:shd"))
    if existing is not None:
        p_pr.remove(existing)
    p_pr.insert(0, _make_shading_element(fill))


def set_style_paragraph_shading(style, fill: str) -> None:
    """Paragraph shading on a document style (e.g. Mark2word Code)."""
    p_pr = style._element.get_or_add_pPr()
    existing = p_pr.find(qn("w:shd"))
    if existing is not None:
        p_pr.remove(existing)
    p_pr.insert(0, _make_shading_element(fill))


def prepare_dual_align_paragraph(
    paragraph,
    content_width_twips: int,
    *,
    suppress_positions: tuple[int, ...] = (),
) -> None:
    """Single right-aligned tab at content width; optional  inherited style tab stops."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = paragraph_element(paragraph).get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is not None:
        p_pr.remove(tabs)
    tabs_el = OxmlElement("w:tabs")
    for pos in suppress_positions:
        clear_tab = OxmlElement("w:tab")
        clear_tab.set(qn("w:val"), "clear")
        clear_tab.set(qn("w:pos"), str(int(pos)))
        tabs_el.append(clear_tab)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "none")
    tab.set(qn("w:pos"), str(int(content_width_twips)))
    tabs_el.append(tab)
    p_pr.append(tabs_el)


def add_dual_align_tab(paragraph) -> None:
    """Insert a tab character after dual-align prep (body or header/footer)."""
    paragraph.add_run().add_tab()


def set_paragraph_border_bottom(paragraph, spec: dict[str, Any]) -> None:
    sz = max(1, round(float(str(spec.get("size", "0.5pt")).rstrip("pt")) * 8))
    color = str(spec.get("color", "000000")).lstrip("#")
    p_pr = paragraph_element(paragraph).get_or_add_pPr()
    p_pr.append(parse_xml(
        f"<w:pBdr {nsdecls('w')}>"
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="2" w:color="{color}"/>'
        f"</w:pBdr>"
    ))


def add_bookmark(paragraph, bookmark_id: int, name: str) -> None:
    p = paragraph_element(paragraph)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, start)
    p.append(end)


def add_page_break(paragraph) -> None:
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_field(paragraph, field_code: str) -> None:
    """Insert a Word field (e.g. PAGE, NUMPAGES)."""
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.extend([fld_begin, instr, fld_sep, fld_end])


def set_run_shading(run, fill: str) -> None:
    """Highlight-style background on a run (inline code)."""
    r_pr = run._r.get_or_add_rPr()
    existing = r_pr.find(qn("w:shd"))
    if existing is not None:
        r_pr.remove(existing)
    r_pr.insert(0, _make_shading_element(fill))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell_element(cell).get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), str(fill).lstrip("#"))
    tc_pr.append(shading)


def set_cell_margins(cell, padding: dict[str, Any]) -> None:
    tc_pr = cell_element(cell).get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for edge, key in (("top", "top"), ("bottom", "bottom"), ("left", "left"), ("right", "right")):
        if key not in padding:
            continue
        pt = _pt_twips(padding[key])
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(pt))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    if len(tc_mar):
        tc_pr.append(tc_mar)


def clear_style_tab_stops(doc, *style_names: str) -> None:
    """Remove built-in Header/Footer tab stops that add to paragraph tabs in Word."""
    for name in style_names:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        p_pr = style._element.find(qn("w:pPr"))
        if p_pr is None:
            continue
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is not None:
            p_pr.remove(tabs)


def clear_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def set_table_layout(table, *, indent_twips: int, width_twips: int) -> None:
    """Fixed-width table indented from the left margin (dxa twips)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for tag in ("w:tblInd", "w:tblW", "w:tblLayout"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(max(0, int(indent_twips))))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(max(0, int(width_twips))))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)


def set_table_width_full(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblW"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)


def set_cell_borders(cell, spec: dict[str, Any]) -> None:
    """Per-edge cell borders; edges not in *spec* (or ``none``) are nil."""
    tc_pr = cell_element(cell).get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        cfg = spec.get(edge)
        el = OxmlElement(f"w:{edge}")
        if isinstance(cfg, dict):
            sz = max(1, round(float(str(cfg.get("size", "0.5pt")).rstrip("pt")) * 8))
            color = str(cfg.get("color", "000000")).lstrip("#")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:color"), color)
            el.set(qn("w:space"), "0")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tc_pr.append(borders)


def set_table_borders(table, spec: dict[str, Any]) -> None:
    sz = max(1, round(float(str(spec.get("size", "0.5pt")).rstrip("pt")) * 8))
    color = str(spec.get("color", "000000")).lstrip("#")
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def set_image_alt(run, alt: str) -> None:
    """Set Word accessibility alt text on an inline picture (wp:docPr + pic:cNvPr)."""
    drawing = run._r.find(f".//{{{W_NS}}}drawing")
    if drawing is None:
        return
    for tag in (f"{{{WP_NS}}}docPr", f"{{{PIC_NS}}}cNvPr"):
        node = drawing.find(f".//{tag}")
        if node is not None:
            node.set("descr", alt)
            node.set("title", alt)


def _pt_twips(value) -> int:
    if isinstance(value, (int, float)):
        pt = float(value)
    else:
        s = str(value).strip().lower()
        if s.endswith("pt"):
            pt = float(s[:-2])
        elif s.endswith("in"):
            pt = float(s[:-2]) * 72.0
        else:
            pt = float(s)
    return int(round(pt * 20))
