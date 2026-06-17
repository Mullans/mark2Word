"""Word document emission."""

from __future__ import annotations

import warnings
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor, Twips

from mark2word.ast import Block, InlineRun
from mark2word.lists import apply_list_numbering
from mark2word.parser import parse_inline, parse_to_ast, split_dual_align
from mark2word.plugins import block_emitter
from mark2word.theme import (
    HEADINGS,
    PAGE_SIZES,
    Resolver,
    deep_merge,
    hex_color,
    pt_value,
    to_length,
)

ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def apply_paragraph_style(paragraph, style: dict[str, Any]) -> None:
    pf = paragraph.paragraph_format
    if style.get("align") in ALIGN:
        paragraph.alignment = ALIGN[style["align"]]
    line = style.get("line")
    if line is not None:
        if isinstance(line, (int, float)):
            pf.line_spacing = float(line)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        else:
            pf.line_spacing = to_length(line)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if style.get("indent_left") is not None:
        pf.left_indent = to_length(style["indent_left"])
    if style.get("indent_first_line") is not None:
        pf.first_line_indent = to_length(style["indent_first_line"])
    if style.get("indent_hanging") is not None:
        pf.first_line_indent = -to_length(style["indent_hanging"])
    if style.get("border_bottom"):
        set_border_bottom(paragraph, style["border_bottom"])


def apply_run_style(run, style: dict[str, Any]) -> None:
    if style.get("font"):
        run.font.name = style["font"]
    if style.get("size") is not None:
        run.font.size = Pt(style["size"])
    if style.get("color"):
        run.font.color.rgb = hex_color(style["color"])
    if style.get("bold") is not None:
        run.bold = style["bold"]
    if style.get("italic") is not None:
        run.italic = style["italic"]


def add_hyperlink(paragraph, text: str, url: str, style: dict[str, Any]) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if style.get("color"):
        color = OxmlElement("w:color")
        color.set(qn("w:val"), str(style["color"]).lstrip("#"))
        r_pr.append(color)
    if style.get("underline") is not False:
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_runs(paragraph, segment: str, style: dict[str, Any], *, code_style: dict[str, Any] | None = None) -> None:
    for run in parse_inline(segment):
        active = code_style if run.code and code_style else style
        if run.url:
            add_hyperlink(paragraph, run.text, run.url, active)
            continue
        r = paragraph.add_run(run.text)
        apply_run_style(r, active)
        if run.bold:
            r.bold = True
        if run.italic:
            r.italic = True


def set_border_bottom(paragraph, spec: dict[str, Any]) -> None:
    sz = max(1, round(float(str(spec.get("size", "0.5pt")).rstrip("pt")) * 8))
    color = str(spec.get("color", "000000")).lstrip("#")
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(parse_xml(
        f"<w:pBdr {nsdecls('w')}>"
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="2" w:color="{color}"/>'
        f"</w:pBdr>"
    ))


def configure_document_styles(doc: Document, resolver: Resolver) -> None:
    base = resolver.resolve("body")
    normal = doc.styles["Normal"]
    normal.font.name = base.get("font", "Calibri")
    normal.font.size = Pt(base.get("size", 11))
    if base.get("color"):
        normal.font.color.rgb = hex_color(base["color"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for level in range(1, 7):
        element = f"h{level}"
        style = resolver.resolve(element)
        heading = doc.styles[f"Heading {level}"]
        if style.get("font"):
            heading.font.name = style["font"]
        if style.get("size") is not None:
            heading.font.size = Pt(style["size"])
        if style.get("color"):
            heading.font.color.rgb = hex_color(style["color"])
        if style.get("bold") is not None:
            heading.font.bold = style["bold"]
        if style.get("italic") is not None:
            heading.font.italic = style["italic"]


def _heading_style_name(block_type: str) -> str | None:
    if block_type in HEADINGS:
        return f"Heading {block_type[1]}"
    return None


def _warn_list_number_mismatch(expected: int, actual: int, line_no: int | None) -> None:
    where = f"line {line_no}" if line_no else "document"
    warnings.warn(
        f"ordered list number {actual} at {where} ignored; continuing from {expected}",
        stacklevel=2,
    )


def emit_paragraph(
    doc,
    block: dict[str, Any],
    style: dict[str, Any],
    content_width_twips: int,
    *,
    prev_list_paragraph=None,
    list_run_num_id: int | None = None,
    expected_ol_number: int | None = None,
    verbose: bool = False,
    md_path: Path | None = None,
    code_style: dict[str, Any] | None = None,
):
    block_type = block["type"]
    custom = block_emitter(block_type)
    if custom is not None:
        p = custom(doc, block, style, {
            "content_width_twips": content_width_twips,
            "md_path": md_path,
            "code_style": code_style,
        })
        return p, None

    list_num_id = None
    if block_type == "list":
        ordered = block.get("ordered", False)
        p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        level = int(block.get("list_level", 0))
        start = block.get("list_number") if ordered else None
        if ordered and expected_ol_number is not None and start is not None and start != expected_ol_number:
            if verbose:
                _warn_list_number_mismatch(expected_ol_number, start, block.get("line_no"))
        list_num_id = apply_list_numbering(
            p,
            doc,
            ordered=ordered,
            start=start,
            level=level,
            prev=prev_list_paragraph,
            run_num_id=list_run_num_id,
        )
    elif block_type == "code":
        p = doc.add_paragraph()
    elif heading := _heading_style_name(block_type):
        p = doc.add_paragraph(style=heading)
    else:
        p = doc.add_paragraph()

    apply_paragraph_style(p, style)
    pf = p.paragraph_format

    if block_type == "image":
        image_path = Path(block["image_path"])
        if not image_path.is_absolute() and md_path is not None:
            image_path = md_path.parent / image_path
        run = p.add_run()
        run.add_picture(str(image_path))
        return p, None

    if block_type == "code":
        for line in block.get("text", "").splitlines() or [""]:
            r = p.add_run(line + "\n")
            apply_run_style(r, code_style or style)
        return p, None

    parts = split_dual_align(block.get("text", ""))
    if parts is not None:
        left, right = parts
        pf.tab_stops.add_tab_stop(Twips(content_width_twips), WD_TAB_ALIGNMENT.RIGHT)
        add_runs(p, left, style, code_style=code_style)
        p.add_run().add_tab()
        add_runs(p, right, style, code_style=code_style)
    else:
        add_runs(p, block.get("text", ""), style, code_style=code_style)
    return p, list_num_id


def apply_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), str(fill).lstrip("#"))
    tc_pr.append(shading)


def emit_table(doc, block: dict[str, Any], resolver: Resolver) -> None:
    rows = block.get("rows", [])
    if not rows:
        return
    region = tuple(block.get("region", []))
    table_style = resolver.resolve("table", region)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            element = "th" if r_idx == 0 else "td"
            cell_style = resolver.resolve(element, region)
            merged = {**table_style, **cell_style}
            if merged.get("fill"):
                apply_cell_shading(cell, merged["fill"])
            p = cell.paragraphs[0]
            add_runs(p, cell_text, merged)


def realize_spacing(doc_paragraphs) -> None:
    for idx, (p, blk, st) in enumerate(doc_paragraphs):
        before = pt_value(st.get("space_before", 0))
        if idx == 0:
            p.paragraph_format.space_before = Pt(before)
            continue
        up_p, up_blk, up_st = doc_paragraphs[idx - 1]
        same = blk.get("type") == up_blk.get("type") == "list"
        if same and ("space_between" in up_st or "space_between" in st):
            gap = pt_value(st.get("space_between", up_st.get("space_between", 0)))
        else:
            gap = max(pt_value(up_st.get("space_after", 0)), before)
        up_p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(gap)
    if doc_paragraphs:
        last_p, _, last_st = doc_paragraphs[-1]
        last_p.paragraph_format.space_after = Pt(pt_value(last_st.get("space_after", 0)))


def emit_from_ast(
    doc: Document,
    blocks: list[Block],
    resolver: Resolver,
    *,
    content_width_twips: int,
    verbose: bool = False,
    md_path: Path | None = None,
) -> list[tuple[Any, dict[str, Any], dict[str, Any]]]:
    emitted = []
    prev_list_p = None
    prev_list_ordered: bool | None = None
    list_run_num_id: int | None = None
    ol_expected: dict[int, int] = {}
    code_style = resolver.resolve("code")

    for block in blocks:
        data = block.to_dict()
        block_type = data["type"]
        if block_type == "table":
            emit_table(doc, data, resolver)
            prev_list_p = None
            prev_list_ordered = None
            list_run_num_id = None
            ol_expected.clear()
            continue

        region = tuple(data.get("region", []))
        if block_type == "list":
            style = resolver.resolve_list_style(int(data.get("list_level", 0)), region)
        else:
            style = resolver.resolve(block_type, region)
        expected = None
        is_ordered = bool(data.get("ordered"))
        same_list_run = (
            block_type == "list"
            and prev_list_p is not None
            and prev_list_ordered == is_ordered
        )
        run_num_id = list_run_num_id if same_list_run else None

        if block_type == "list" and is_ordered:
            level = int(data.get("list_level", 0))
            if not same_list_run:
                ol_expected.clear()
            if level not in ol_expected:
                ol_expected[level] = data.get("list_number") or 1
            expected = ol_expected[level]
            ol_expected[level] = expected + 1

        list_prev = prev_list_p if same_list_run else None

        p, new_num_id = emit_paragraph(
            doc,
            data,
            style,
            content_width_twips,
            prev_list_paragraph=list_prev,
            list_run_num_id=run_num_id,
            expected_ol_number=expected,
            verbose=verbose,
            md_path=md_path,
            code_style=code_style,
        )
        if block_type == "list":
            prev_list_p = p
            prev_list_ordered = is_ordered
            list_run_num_id = new_num_id
        else:
            prev_list_p = None
            prev_list_ordered = None
            list_run_num_id = None
            ol_expected.clear()
        emitted.append((p, data, style))
    return emitted


def build(
    glob: dict[str, Any],
    external: dict[str, Any],
    md_body: str,
    out_path: Path,
    *,
    verbose: bool = False,
    md_path: Path | None = None,
) -> None:
    resolver = Resolver(external, glob)
    doc = Document()

    page = deep_merge(external.get("page", {}), glob.get("page", {}))
    sec = doc.sections[0]
    size_key = str(page.get("size", "letter")).lower()
    if size_key not in PAGE_SIZES:
        raise ValueError(f"Unknown page size: `{page.get('size')}`")
    sec.page_width, sec.page_height = PAGE_SIZES[size_key]
    margin = page.get("margin", {})
    sec.top_margin = to_length(margin.get("top", "0.5in"))
    sec.bottom_margin = to_length(margin.get("bottom", "0.5in"))
    sec.left_margin = to_length(margin.get("left", "0.7in"))
    sec.right_margin = to_length(margin.get("right", "0.7in"))
    content_emu = int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)
    content_width_twips = int(content_emu / 635)

    configure_document_styles(doc, resolver)
    blocks = parse_to_ast(md_body)
    emitted = emit_from_ast(
        doc,
        blocks,
        resolver,
        content_width_twips=content_width_twips,
        verbose=verbose,
        md_path=md_path,
    )
    realize_spacing(emitted)
    doc.save(str(out_path))
    if verbose:
        print(f"Wrote {out_path}")
    else:
        print(f"Wrote {out_path}")
