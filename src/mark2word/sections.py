"""Document header/footer from theme page settings."""

from __future__ import annotations

import re
from typing import Any

from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Twips

from mark2word.oxml_helpers import add_field, add_run_text, paragraph_element
from mark2word.parser import split_dual_align

_PLACEHOLDER = re.compile(r"\{page\}|\{pages\}|\{title\}")


def configure_page_chrome(
    doc,
    page: dict[str, Any],
    *,
    doc_title: str,
    content_width_twips: int,
) -> None:
    header = page.get("header")
    footer = page.get("footer")
    if not header and not footer:
        return
    sec = doc.sections[0]
    meta = {"title": doc_title}
    if header:
        _apply_chrome_slot(sec.header, header, meta, content_width_twips)
    if footer:
        _apply_chrome_slot(sec.footer, footer, meta, content_width_twips)


def _apply_chrome_slot(container, cfg, meta: dict[str, str], content_width_twips: int) -> None:
    text = cfg if isinstance(cfg, str) else str(cfg.get("text", ""))
    if not text.strip():
        return
    p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    p.text = ""
    parts = split_dual_align(text)
    if parts:
        left, right = parts
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(Twips(content_width_twips), WD_TAB_ALIGNMENT.RIGHT)
        _render_chrome_segment(p, left, meta)
        p.add_run().add_tab()
        _render_chrome_segment(p, right, meta)
    else:
        _render_chrome_segment(p, text, meta)


def _render_chrome_segment(paragraph, text: str, meta: dict[str, str]) -> None:
    pos = 0
    p_el = paragraph_element(paragraph)
    for match in _PLACEHOLDER.finditer(text):
        if match.start() > pos:
            add_run_text(p_el, text[pos:match.start()])
        token = match.group(0)
        if token == "{page}":
            add_field(paragraph, "PAGE")
        elif token == "{pages}":
            add_field(paragraph, "NUMPAGES")
        elif token == "{title}":
            add_run_text(p_el, meta.get("title", ""))
        pos = match.end()
    if pos < len(text):
        add_run_text(p_el, text[pos:])
